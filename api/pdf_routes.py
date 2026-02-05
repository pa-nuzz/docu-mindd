
#pdf_routes
from fastapi import APIRouter, UploadFile, File, HTTPException
from pypdf import PdfReader
from docx import Document
from pathlib import Path
import io

from services.ocr_utils import extract_text_with_ocr
from services.text_cleaning import clean_text

router = APIRouter(prefix="/upload", tags=["Documents"])

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
UPLOAD_FOLDER = Path("uploads")
UPLOAD_FOLDER.mkdir(exist_ok=True)


@router.post("/file")
async def upload_file(file: UploadFile = File(...)):
    filename = file.filename.lower()

    # Only PDF or DOCX
    if not (filename.endswith(".pdf") or filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only PDF or DOCX allowed")

    # Read file content
    content = await file.read()

    # Check size
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large")

    # Save the file
    file_path = UPLOAD_FOLDER / file.filename
    with open(file_path, "wb") as f:
        f.write(content)

    # Extract text
    raw_text = ""
    page_count = 0
    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(content))
        raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        page_count = len(reader.pages)
    else:
        doc = Document(io.BytesIO(content))
        raw_text = "\n".join(para.text for para in doc.paragraphs)
        page_count = len(doc.paragraphs)

    # OCR fallback if text is empty
    if not raw_text.strip():
        raw_text = extract_text_with_ocr(content)

    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="No readable text found")

    # Clean the text
    cleaned_text = clean_text(raw_text)

    return {
        "filename": file.filename,
        "page_count": page_count,
        "clean_text": cleaned_text
    }
